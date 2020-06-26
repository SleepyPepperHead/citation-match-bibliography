import re
import docx
from docx.shared import RGBColor


text = '''
main part of the article
'''
# citation in paren
regex_p = re.compile('\(.*?\)')
# citation not in paren
regex = re.compile('.{20}\(\d{4}\w?\)')
no_paren = regex.findall(text)
in_paren = regex_p.findall(text)

# split all stuff in paren adn add to the list raw
raw = []
for p in in_paren:
    p = p.strip('()')
    raw.extend(p.split(';'))

# filter (author_year) items to the list raw_pro
raw_pro = []
for i in raw:
    if ('20' in i or '19' in i) and len(i) > 5:
        i = i.strip('e.g.,')
        i = i.strip()
        i = i.strip('and references therei')
        raw_pro.append(i)

print(len(raw_pro))
# print no_paren and add manually
add = '''Luo et al. 2015;Huang et al. 2016;Yang et al. 2015;Henry and Guidotti 1985;Yang and Jiang 2012;Yang and 
Jiang 2002;Henry and Dutrow 2012;Yardley 2005'''
raw_pro.extend(add.split(';'))
# print to check and add double years manually
add2 = "Qiu 2016;Yang 2015a;Yang 2015b;Deng 2017;Yu 2020a;Yu 2020b"
raw_pro.extend(add2.split(';'))
refers = set(raw_pro)
print(len(raw_pro))
print(len(refers))

a_y = []
for refer in refers:
    s = refer.split()
    a_y.append((s[0], s[-1]))

# drop duplicated in (author, year) style
ay = list(set(a_y))
ay.sort()
print(ay)

doc = docx.Document('refer.docx')
lines = [para.text for para in doc.paragraphs]
for author in ay:
    for para in doc.paragraphs:
        i = para.text
        if author[0] in i[0:11] and author[-1] in i:
            # right paragraph will be marked red
            for run in para.runs:
                run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
            break
        if lines[-1] == i:
            print(author[0], author[1], 'not in reference')
doc.save("new_refer_2.docx")



